"""Serveur web : upload audio, transcription streamée découplée du SSE, renommage des speakers, téléchargement.

Le worker tourne dans un thread de fond et écrit les fichiers .txt/.srt au fur et à mesure :
si le client se déconnecte (PC verrouillé, onglet fermé, crash du navigateur), la transcription
continue côté serveur et tout ce qui a déjà été transcrit est sur disque.

Les events sont bufferisés par job, donc une reconnexion SSE rejoue tout l'historique avant
de basculer sur le live."""
import io
import json
import os
import queue
import shutil
import threading
import uuid
import zipfile
from pathlib import Path
from typing import Optional

import soundfile as sf
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles

from core import (
    extract_audio_segment,
    format_timestamp,
    load_diarization_pipeline,
    load_whisper,
    transcribe_stream,
    write_outputs,
)


def _ts_hms(seconds: float) -> str:
    return format_timestamp(seconds).split(",")[0]


def build_markdown(job: dict) -> str:
    name_map = job.get("name_map") or {}
    lines = [f"# {job['stem']}", ""]
    speakers_in_use = {name_map.get(seg.get("speaker"), seg.get("speaker")) for seg in job["segments"] if seg.get("speaker")}
    if speakers_in_use:
        lines.append("**Locuteurs** : " + ", ".join(sorted(speakers_in_use)))
        lines.append("")
    current_speaker = None
    for seg in job["segments"]:
        speaker = seg.get("speaker")
        display = name_map.get(speaker, speaker) if speaker else None
        ts = _ts_hms(seg["start"])
        if display:
            if display != current_speaker:
                lines.append("")
                lines.append(f"### {display}")
                lines.append("")
                current_speaker = display
            lines.append(f"_{ts}_ — {seg['text']}")
        else:
            lines.append(f"_{ts}_ — {seg['text']}")
    return "\n".join(lines) + "\n"


def build_docx(job: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    name_map = job.get("name_map") or {}
    doc = Document()
    doc.add_heading(job["stem"], level=1)
    speakers_in_use = sorted({name_map.get(seg.get("speaker"), seg.get("speaker")) for seg in job["segments"] if seg.get("speaker")})
    if speakers_in_use:
        p = doc.add_paragraph()
        r = p.add_run("Locuteurs : "); r.bold = True
        p.add_run(", ".join(speakers_in_use))
    doc.add_paragraph()

    pink = RGBColor(0xDB, 0x27, 0x77)
    gray = RGBColor(0x99, 0x99, 0x99)
    current_speaker = None
    for seg in job["segments"]:
        speaker = seg.get("speaker")
        display = name_map.get(speaker, speaker) if speaker else None
        ts = _ts_hms(seg["start"])
        if display and display != current_speaker:
            p = doc.add_paragraph()
            r = p.add_run(display); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = pink
            current_speaker = display
        p = doc.add_paragraph()
        ts_r = p.add_run(f"{ts}  "); ts_r.italic = True; ts_r.font.size = Pt(9); ts_r.font.color.rgb = gray
        p.add_run(seg["text"])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

load_dotenv()

UPLOADS_DIR = Path("uploads")
OUTPUT_DIR = Path("output")
STATIC_DIR = Path("static")
UPLOADS_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

_WHISPER_CACHE: dict = {}
_DIAR_PIPELINE = {"pipeline": None}
JOBS: dict = {}

app = FastAPI(title="Clo — Speech to Text")
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


@app.get("/", response_class=HTMLResponse)
def index():
    return STATIC_DIR.joinpath("index.html").read_text(encoding="utf-8")


@app.get("/favicon.svg")
def favicon():
    return FileResponse(STATIC_DIR / "favicon.svg", media_type="image/svg+xml")


@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    job_id = uuid.uuid4().hex[:12]
    src_name = file.filename or "audio.mp3"
    ext = Path(src_name).suffix.lower() or ".mp3"
    stem = Path(src_name).stem or "audio"
    audio_path = UPLOADS_DIR / f"{job_id}{ext}"
    with audio_path.open("wb") as f:
        shutil.copyfileobj(file.file, f)
    JOBS[job_id] = {
        "audio_path": audio_path,
        "stem": stem,
        "filename": src_name,
        "out_dir": OUTPUT_DIR / stem,
        "segments": [],
        "speakers": [],
        "name_map": {},
        "status": "pending",  # pending | running | done | error
        "events": [],
        "subscribers": [],
        "lock": threading.Lock(),
        "thread": None,
        "error": None,
        "config": {},
    }
    return {"job_id": job_id, "filename": src_name}


def _get_whisper(model: str):
    key = (model, "auto", "int8")
    if key not in _WHISPER_CACHE:
        _WHISPER_CACHE[key] = load_whisper(model)
    return _WHISPER_CACHE[key]


def _get_diar_pipeline():
    if _DIAR_PIPELINE["pipeline"] is None:
        token = os.environ.get("HF_TOKEN") or os.environ.get("HUGGINGFACE_TOKEN")
        if not token:
            raise RuntimeError("HF_TOKEN manquant — impossible de lancer la diarisation.")
        _DIAR_PIPELINE["pipeline"] = load_diarization_pipeline(token)
    return _DIAR_PIPELINE["pipeline"]


def _publish(job: dict, event: dict) -> None:
    with job["lock"]:
        job["events"].append(event)
        subs = list(job["subscribers"])
    for q in subs:
        try:
            q.put_nowait(event)
        except queue.Full:
            pass


def _run_job(job_id: str, model: str, language: Optional[str], diarize: bool, vad: bool) -> None:
    """Worker exécuté dans un thread. Écrit les fichiers au fur et à mesure."""
    job = JOBS[job_id]
    job["status"] = "running"
    job["config"] = {"model": model, "language": language, "diarize": diarize, "vad": vad}

    out_dir = job["out_dir"]
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = job["stem"]
    txt_path = out_dir / f"{stem}.txt"
    srt_path = out_dir / f"{stem}.srt"
    txt_file = None
    srt_file = None
    current_speaker_in_txt = None

    try:
        whisper_model = _get_whisper(model)
        diar_pipeline = _get_diar_pipeline() if diarize else None

        txt_file = txt_path.open("w", encoding="utf-8")
        srt_file = srt_path.open("w", encoding="utf-8")

        lang = language if language and language != "auto" else None
        for event in transcribe_stream(
            job["audio_path"],
            whisper_model,
            diar_pipeline=diar_pipeline,
            language=lang,
            vad=vad,
        ):
            if event["type"] == "segment":
                text = event["text"]
                speaker = event.get("speaker")
                if speaker:
                    if speaker != current_speaker_in_txt:
                        if current_speaker_in_txt is not None:
                            txt_file.write("\n")
                        txt_file.write(f"{speaker}:\n")
                        current_speaker_in_txt = speaker
                    txt_file.write(f"  {text}\n")
                    srt_text = f"[{speaker}] {text}"
                else:
                    txt_file.write(text + "\n")
                    srt_text = text
                srt_file.write(f"{event['index']}\n")
                srt_file.write(f"{format_timestamp(event['start'])} --> {format_timestamp(event['end'])}\n")
                srt_file.write(f"{srt_text}\n\n")
                txt_file.flush()
                srt_file.flush()
                job["segments"].append({
                    "index": event["index"],
                    "start": event["start"],
                    "end": event["end"],
                    "text": event["text"],
                    "speaker": event.get("speaker"),
                })
            elif event["type"] == "done":
                job["speakers"] = event["speakers"]
                for spk in event["speakers"]:
                    spk["sample_url"] = f"/sample/{job_id}/{spk['label'].replace(' ', '_')}.wav"
                event["downloads"] = {
                    "txt": f"/download/{job_id}/txt",
                    "srt": f"/download/{job_id}/srt",
                    "zip": f"/download/{job_id}/zip",
                }
            _publish(job, event)

        job["status"] = "done"
    except Exception as e:
        job["status"] = "error"
        job["error"] = str(e)
        _publish(job, {"type": "error", "message": str(e)})
    finally:
        if txt_file:
            txt_file.close()
        if srt_file:
            srt_file.close()


@app.post("/start/{job_id}")
def start_job(
    job_id: str,
    model: str = "small",
    language: Optional[str] = "fr",
    diarize: bool = False,
    vad: bool = True,
):
    if job_id not in JOBS:
        raise HTTPException(404, "Job inconnu")
    job = JOBS[job_id]
    if job["status"] in ("running", "done"):
        return {"ok": True, "status": job["status"]}
    t = threading.Thread(
        target=_run_job, args=(job_id, model, language, diarize, vad), daemon=True
    )
    job["thread"] = t
    t.start()
    return {"ok": True, "status": "running"}


@app.get("/events/{job_id}")
def events(job_id: str):
    """SSE qui rejoue d'abord tout l'historique du job puis suit le live."""
    if job_id not in JOBS:
        raise HTTPException(404, "Job inconnu")
    job = JOBS[job_id]

    q: queue.Queue = queue.Queue(maxsize=10000)
    with job["lock"]:
        for ev in job["events"]:
            try:
                q.put_nowait(ev)
            except queue.Full:
                pass
        if job["status"] in ("pending", "running"):
            job["subscribers"].append(q)

    def stream():
        try:
            while True:
                try:
                    ev = q.get(timeout=15)
                except queue.Empty:
                    yield ": keepalive\n\n"
                    if job["status"] in ("done", "error") and q.empty():
                        return
                    continue
                yield "data: " + json.dumps(ev, ensure_ascii=False) + "\n\n"
                if ev["type"] in ("done", "error"):
                    return
        finally:
            with job["lock"]:
                try:
                    job["subscribers"].remove(q)
                except ValueError:
                    pass

    return StreamingResponse(stream(), media_type="text/event-stream")


@app.get("/job/{job_id}")
def job_status(job_id: str):
    """Permet au frontend de vérifier l'état d'un job (pour reprise après crash navigateur)."""
    if job_id not in JOBS:
        raise HTTPException(404, "Job inconnu")
    job = JOBS[job_id]
    return {
        "job_id": job_id,
        "filename": job["filename"],
        "status": job["status"],
        "error": job["error"],
        "n_segments": len(job["segments"]),
        "n_speakers": len(job["speakers"]),
        "config": job.get("config", {}),
    }


@app.get("/sample/{job_id}/{filename}")
def sample(job_id: str, filename: str):
    if job_id not in JOBS:
        raise HTTPException(404)
    job = JOBS[job_id]
    label = Path(filename).stem.replace("_", " ")
    for spk in job["speakers"]:
        if spk["label"] == label:
            audio = extract_audio_segment(job["audio_path"], spk["sample_start"], spk["sample_end"])
            buf = io.BytesIO()
            sf.write(buf, audio, 16000, format="WAV")
            buf.seek(0)
            return Response(content=buf.read(), media_type="audio/wav")
    raise HTTPException(404, "Speaker introuvable")


@app.post("/rename/{job_id}")
async def rename(job_id: str, request: Request):
    if job_id not in JOBS:
        raise HTTPException(404)
    job = JOBS[job_id]
    body = await request.json()
    name_map = {str(k): str(v).strip() for k, v in body.items() if v and str(v).strip()}
    job["name_map"] = name_map
    write_outputs(job["out_dir"], job["stem"], job["segments"], name_map=name_map)
    return {"ok": True, "name_map": name_map}


@app.post("/edit/{job_id}")
async def edit_segment(job_id: str, request: Request):
    """Modifie le texte d'un segment et réécrit les fichiers."""
    if job_id not in JOBS:
        raise HTTPException(404)
    job = JOBS[job_id]
    body = await request.json()
    idx = int(body["index"])
    new_text = str(body.get("text", "")).strip()
    if not (1 <= idx <= len(job["segments"])):
        raise HTTPException(400, "Index invalide")
    if not new_text:
        raise HTTPException(400, "Texte vide")
    job["segments"][idx - 1]["text"] = new_text
    write_outputs(job["out_dir"], job["stem"], job["segments"], name_map=job.get("name_map") or {})
    return {"ok": True}


@app.get("/audio/{job_id}")
def audio(job_id: str):
    """Sert le fichier audio uploadé pour permettre la lecture côté navigateur."""
    if job_id not in JOBS:
        raise HTTPException(404)
    job = JOBS[job_id]
    if not job["audio_path"].exists():
        raise HTTPException(404, "Audio introuvable")
    return FileResponse(job["audio_path"])


@app.post("/reassign/{job_id}")
async def reassign(job_id: str, request: Request):
    """Réassigne manuellement un segment à un autre locuteur (ou aucun)."""
    if job_id not in JOBS:
        raise HTTPException(404)
    job = JOBS[job_id]
    body = await request.json()
    idx = int(body["index"])
    speaker = body.get("speaker")
    if speaker == "":
        speaker = None
    if not (1 <= idx <= len(job["segments"])):
        raise HTTPException(400, "Index de segment invalide")
    job["segments"][idx - 1]["speaker"] = speaker
    if speaker and not any(s["label"] == speaker for s in job["speakers"]):
        job["speakers"].append({"label": speaker, "sample_start": 0.0, "sample_end": 0.0})
    write_outputs(job["out_dir"], job["stem"], job["segments"], name_map=job.get("name_map") or {})
    return {"ok": True}


@app.get("/download/{job_id}/{kind}")
def download(job_id: str, kind: str):
    if job_id not in JOBS:
        raise HTTPException(404)
    job = JOBS[job_id]
    stem = job["stem"]
    txt_path = job["out_dir"] / f"{stem}.txt"
    srt_path = job["out_dir"] / f"{stem}.srt"

    if kind == "txt":
        return FileResponse(txt_path, filename=f"{stem}.txt", media_type="text/plain")
    if kind == "srt":
        return FileResponse(srt_path, filename=f"{stem}.srt", media_type="application/x-subrip")
    if kind == "md":
        return Response(
            content=build_markdown(job).encode("utf-8"),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{stem}.md"'},
        )
    if kind == "docx":
        return Response(
            content=build_docx(job),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{stem}.docx"'},
        )
    if kind == "zip":
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(txt_path, arcname=f"{stem}.txt")
            zf.write(srt_path, arcname=f"{stem}.srt")
        buf.seek(0)
        return Response(
            content=buf.read(),
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{stem}.zip"'},
        )
    raise HTTPException(400, "Kind inconnu")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
